import os
import zipfile
import tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging
import re
import unicodedata

logger = logging.getLogger(__name__)

def add_formatted_text(paragraph, text, formatting_info=None, font_size=8, bold=False):
    """
    Adds text to a paragraph with formatting (subscript/superscript) if available.
    
    Args:
        paragraph: The paragraph object to add text to
        text: The plain text (used if no formatting_info)
        formatting_info: List of tuples (text, is_subscript, is_superscript) or None
        font_size: Font size in points
        bold: Whether to make text bold
    """
    if formatting_info and len(formatting_info) > 0:
        # Use rich text formatting
        for text_part, is_subscript, is_superscript in formatting_info:
            run = paragraph.add_run(text_part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            run.bold = bold
            
            # Apply subscript or superscript
            if is_subscript:
                run.font.subscript = True
            elif is_superscript:
                run.font.superscript = True
    else:
        # Plain text
        run = paragraph.add_run(str(text))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.bold = bold
    
    return paragraph

def create_word_document(questions_df, highlight_answers=False, class_name="", subject_name="", version=0, num_columns=2):
    """
    Creates a Word document with the given questions
    If highlight_answers is True, the correct answers are highlighted
    Optional class_name and subject_name parameters for document header
    num_columns: Number of columns (1 or 2, default is 2)
    """
    doc = Document()
    
    section = doc.sections[0]
    
    # Set orientation and margins based on number of columns
    # Use A4 portrait page size for all documents
    section.orientation = WD_ORIENTATION.PORTRAIT
    # A4 size in inches: 8.27 x 11.69
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # Set margins (slightly narrow but printable)
    # For two-column layouts we keep moderate margins to allow room for columns
    if num_columns == 1:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    else:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Add document header with subject and class name
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create title text
    title_text = "ĐỀ THI"
    if subject_name:
        title_text += f" {subject_name}"
    if class_name:
        title_text += f" - {class_name}"
    
    # Add the title in bold, centered
    run = header.add_run(title_text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    
    # Add version number centered
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run(f"Đề số {version + 1}")  # version starts from 0 in code
    version_run.bold = True
    version_run.font.name = 'Times New Roman'
    version_run.font.size = Pt(8)
    
    # Add student information section without underlines
    info_para = doc.add_paragraph()
    info_para.add_run("Mã sinh viên:                     Họ tên:                    ").font.name = 'Times New Roman'
    info_para.add_run().font.size = Pt(8)
    
    # Add a small space after the student info section
    doc.add_paragraph()
    
    # Set up column layout for questions section
    section.start_type = WD_SECTION.NEW_PAGE
    
    # Create columns based on num_columns parameter
    sectPr = section._sectPr
    if not sectPr.xpath('./w:cols'):
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), str(num_columns))
        sectPr.append(cols)
    else:
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), str(num_columns))
    
    # Try to set font style as a default for the document
    try:
        style = doc.styles['Normal']
        # Set default font and size for the whole document
        style.font.name = 'Times New Roman'
        style.font.size = Pt(8)
        # Reduce paragraph spacing and line height
        para_fmt = style.paragraph_format
        para_fmt.space_before = Pt(0)
        para_fmt.space_after = Pt(1)
        try:
            para_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except Exception:
            # Fallback: set explicit line spacing in points
            para_fmt.line_spacing = Pt(9)
    except Exception as e:
        logger.debug(f"Error accessing style: {str(e)}")
        
    # Process questions
    total_questions = len(questions_df)
    
    for index, row in questions_df.iterrows():
        # Get formatting info if available
        formatting = row.get('_formatting', {})
        
        # Add question number and text (in bold)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)  # Minimal space
        # Justify question paragraph
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add question number
        run = paragraph.add_run(f"{index + 1}. ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        
        # Add question text with formatting
        question_formatting = formatting.get('Câu hỏi', None)
        add_formatted_text(paragraph, row['Câu hỏi'], question_formatting, font_size=8, bold=True)
        
        # Add options with minimal spacing and indentation
        options = ['A', 'B', 'C', 'D']
        correct_answer = row['đáp án']
        
        for option in options:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(8)  # Smaller indentation
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # Justify option paragraph
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add option label
            is_bold = highlight_answers and option == correct_answer
            run = paragraph.add_run(f"{option}: ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
            run.bold = is_bold
            
            # Add option text with formatting
            option_formatting = formatting.get(option, None)
            add_formatted_text(paragraph, row[option], option_formatting, font_size=8, bold=is_bold)
        
        # Add minimal separator between questions
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
    
    return doc

def generate_zip_files(questions_df, num_questions, num_versions, class_name="", subject_name="", random_seed=None, shuffle_answers=False):
    """
    Generates two ZIP files:
    1. Regular version without highlighted answers
    2. Version with highlighted answers
    
    Each version contains both 1-column and 2-column layouts
    
    Args:
        questions_df: DataFrame with questions
        num_questions: Number of questions per version
        num_versions: Number of versions to generate
        class_name: Class name for header
        subject_name: Subject name for header
        random_seed: Optional seed for reproducibility
        shuffle_answers: Whether to shuffle answer options
    
    Returns dictionary with paths to both ZIP files
    """
    temp_dir = tempfile.gettempdir()
    
    # Helper to normalize class name into safe filename
    def _normalize_filename(s: str) -> str:
        if not s:
            return ''
        # Normalize unicode characters to ASCII equivalents
        s = unicodedata.normalize('NFKD', s).encode('ascii', 'ignore').decode('ascii')
        s = s.lower().strip()
        s = re.sub(r'\s+', '_', s)
        s = re.sub(r'[^a-z0-9_-]', '', s)
        return s or ''

    # Create unique filenames for the ZIP files
    regular_zip_filename = f"regular_quiz_{num_questions}q_{num_versions}v.zip"
    highlighted_zip_filename = f"highlighted_quiz_{num_questions}q_{num_versions}v.zip"

    # Full zip filename should be the normalized class name (fallback if empty)
    normalized_class = _normalize_filename(class_name)
    if not normalized_class:
        normalized_class = f"quiz_{num_questions}q_{num_versions}v"

    full_zip_filename = f"{normalized_class}.zip"

    regular_zip_path = os.path.join(temp_dir, regular_zip_filename)
    highlighted_zip_path = os.path.join(temp_dir, highlighted_zip_filename)
    full_zip_path = os.path.join(temp_dir, full_zip_filename)

    # Create all three ZIP files
    with zipfile.ZipFile(regular_zip_path, 'w') as regular_zip, \
         zipfile.ZipFile(highlighted_zip_path, 'w') as highlighted_zip, \
         zipfile.ZipFile(full_zip_path, 'w') as full_zip:

        # We'll collect the per-version question DataFrames to build the answer-key Excel
        per_version_questions = []

        for version in range(1, num_versions + 1):
            # Calculate seed for this version if base seed provided
            version_seed = random_seed + version if random_seed is not None else None

            # Get random questions for this version
            version_questions = get_random_questions(questions_df, num_questions, version_seed)

            # Shuffle answers if requested
            if shuffle_answers:
                version_questions = shuffle_question_answers(version_questions, version_seed)

            # Save a single 2-column document per version (user requested only 2-column)
            doc = create_word_document(
                version_questions,
                highlight_answers=False,
                class_name=class_name,
                subject_name=subject_name,
                version=version-1,
                num_columns=2
            )

            # Normalized docx filename: {normalized_class}_{version}.docx
            doc_filename = f"{normalized_class}_{version}.docx"
            doc_path = os.path.join(temp_dir, doc_filename)
            doc.save(doc_path)

            # Add to regular and highlighted zips for compatibility (both contain the statement docx)
            regular_zip.write(doc_path, doc_filename)
            highlighted_zip.write(doc_path, doc_filename)

            # Add to full zip as well
            full_zip.write(doc_path, doc_filename)

            # Store the DataFrame for the Excel answer-key
            per_version_questions.append((f"V{version}", version_questions.copy()))

            # Clean up temporary docx
            os.remove(doc_path)

        # After generating all versions, build the Excel answer-key with one sheet per version
        try:
            import pandas as pd

            excel_filename = f"{normalized_class}_answer_key.xlsx"
            excel_path = os.path.join(temp_dir, excel_filename)

            with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                for sheet_name, df in per_version_questions:
                    # Build sheet with columns: STT câu, nội dung câu, A, B, C, D, đáp án, phân loại
                    rows = []
                    for idx, row in df.iterrows():
                        stt = idx + 1
                        content = row.get('Câu hỏi', '')
                        a = row.get('A', '')
                        b = row.get('B', '')
                        c = row.get('C', '')
                        d = row.get('D', '')
                        key = row.get('đáp án', '')
                        category = row.get('Phân loại', '') if 'Phân loại' in row.index else ''
                        rows.append({
                            'STT câu': stt,
                            'Nội dung câu': content,
                            'A': a,
                            'B': b,
                            'C': c,
                            'D': d,
                            'đáp án': key,
                            'phân loại': category
                        })

                    sheet_df = pd.DataFrame(rows, columns=['STT câu', 'Nội dung câu', 'A', 'B', 'C', 'D', 'đáp án', 'phân loại'])
                    # Write to a sheet named V1, V2, ...
                    sheet_df.to_excel(writer, sheet_name=sheet_name, index=False)

            # Add the excel file into the full ZIP
            full_zip.write(excel_path, excel_filename)

            # Clean up excel file
            os.remove(excel_path)
        except Exception as e:
            logger.exception(f"Failed to create answer-key Excel: {e}")
            # If Excel creation fails, proceed without adding it
            pass
    
    return {
        'regular': regular_zip_filename,
        'highlighted': highlighted_zip_filename,
        'full': full_zip_filename
    }

def shuffle_question_answers(questions_df, random_seed=None):
    """
    Shuffles the answer options (A, B, C, D) for each question and updates the correct answer.
    
    Args:
        questions_df: DataFrame containing questions
        random_seed: Optional seed for reproducibility
        
    Returns:
        DataFrame with shuffled answers
    """
    import pandas as pd
    import numpy as np
    
    # Create a copy to avoid modifying the original
    shuffled_df = questions_df.copy()
    
    # Set random seed if provided
    if random_seed is not None:
        np.random.seed(random_seed)
    
    # Define the answer options
    options = ['A', 'B', 'C', 'D']
    
    for idx, row in shuffled_df.iterrows():
        # Get current answer values
        answer_values = {opt: row[opt] for opt in options}
        correct_answer = row['đáp án']
        
        # Get formatting info if available
        formatting = row.get('_formatting', {})
        answer_formatting = {opt: formatting.get(opt, None) for opt in options}
        
        # Create a shuffled mapping
        shuffled_options = options.copy()
        np.random.shuffle(shuffled_options)
        
        # Create new mapping: new position -> old value
        new_mapping = {options[i]: answer_values[shuffled_options[i]] for i in range(4)}
        new_formatting_mapping = {options[i]: answer_formatting[shuffled_options[i]] for i in range(4)}
        
        # Update the row with shuffled values
        for opt in options:
            shuffled_df.at[idx, opt] = new_mapping[opt]
        
        # Update formatting for shuffled answers
        if formatting:
            new_formatting = formatting.copy()
            for opt in options:
                if new_formatting_mapping[opt] is not None:
                    new_formatting[opt] = new_formatting_mapping[opt]
            shuffled_df.at[idx, '_formatting'] = new_formatting
        
        # Update the correct answer
        # Find which new position has the correct answer value
        correct_value = answer_values[correct_answer]
        for new_opt, value in new_mapping.items():
            if value == correct_value:
                shuffled_df.at[idx, 'đáp án'] = new_opt
                break
    
    return shuffled_df

def get_random_questions(df, num_questions, random_seed=None):
    """
    Gets a random sample of questions from the dataframe.
    If "Phân loại" column exists, ensures at least 1 question from each category.
    
    Args:
        df: DataFrame containing questions
        num_questions: Number of questions to select
        random_seed: Optional seed for reproducibility
    """
    import pandas as pd
    import numpy as np
    
    # Set random seed if provided
    if random_seed is not None:
        np.random.seed(random_seed)
    
    # Check if "Phân loại" column exists
    if 'Phân loại' not in df.columns:
        # No category column, use simple random selection
        return df.sample(n=num_questions, random_state=random_seed).reset_index(drop=True)
    
    # Filter out rows with null/empty categories
    df_with_categories = df[df['Phân loại'].notna()].copy()
    
    # If no valid categories exist, fall back to simple random selection from all rows
    if len(df_with_categories) == 0:
        logger.warning('Cột "Phân loại" tồn tại nhưng không chứa giá trị hợp lệ nào. Sử dụng random đơn giản.')
        return df.sample(n=num_questions, random_state=random_seed).reset_index(drop=True)
    
    # Get unique categories (excluding NaN)
    categories = df_with_categories['Phân loại'].unique()
    num_categories = len(categories)
    
    # Check if we have enough questions to cover all categories
    if num_questions < num_categories:
        raise ValueError(f'Số câu hỏi ({num_questions}) phải lớn hơn hoặc bằng số phân loại ({num_categories}) để đảm bảo mỗi phân loại có ít nhất 1 câu hỏi')
    
    # Validate each category has at least one question
    empty_categories = []
    for category in categories:
        category_questions = df_with_categories[df_with_categories['Phân loại'] == category]
        if len(category_questions) == 0:
            empty_categories.append(str(category))
    
    if empty_categories:
        raise ValueError(f'Các phân loại sau không có câu hỏi: {", ".join(empty_categories)}')
    
    # Select at least 1 question from each category
    selected_questions = []
    selected_indices = []
    
    for category in categories:
        category_questions = df_with_categories[df_with_categories['Phân loại'] == category]
        # Randomly select 1 question from this category
        sampled = category_questions.sample(n=1, random_state=random_seed)
        selected_questions.append(sampled)
        selected_indices.extend(sampled.index.tolist())
        # Update seed for next iteration if seed is provided
        if random_seed is not None:
            random_seed += 1
    
    # Calculate remaining questions to select
    remaining_questions = num_questions - num_categories
    
    if remaining_questions > 0:
        # Create a pool of remaining questions (all questions except already selected)
        remaining_df = df_with_categories[~df_with_categories.index.isin(selected_indices)]
        
        # Randomly select the remaining questions
        additional_questions = remaining_df.sample(n=remaining_questions, random_state=random_seed)
        
        # Add to selected questions
        selected_questions.append(additional_questions)
        # Update seed for final shuffle if seed is provided
        if random_seed is not None:
            random_seed += 1
    
    # Combine all selected questions
    result_df = pd.concat(selected_questions, ignore_index=True)
    
    # Shuffle the final result to mix categories
    result_df = result_df.sample(frac=1, random_state=random_seed).reset_index(drop=True)
    
    return result_df
